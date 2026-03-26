"""
Word 教案生成引擎

对应 A04 要求:
  - 4b) 生成与PPT配套的详细教案，包括教学目标、教学过程、
        教学方法、课堂活动设计、课后作业等
  - 5c) 支持将最终课件以 .docx 格式下载
"""

from __future__ import annotations

import json
import re
import time
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


# ═══════════════════════════════════════════
# 样式辅助
# ═══════════════════════════════════════════

def _add_heading(doc: Document, text: str, level: int = 1):
    """添加标题"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = "微软雅黑"
    return heading


def _add_para(doc: Document, text: str, bold: bool = False, font_size: int = 12,
              alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after: int = 6):
    """添加段落"""
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.bold = bold
    run.font.name = "微软雅黑"
    return p


def _add_table(doc: Document, headers: List[str], rows: List[List[str]]):
    """添加表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(11)

    # 数据行
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(11)

    doc.add_paragraph()  # 表后空行
    return table


# ═══════════════════════════════════════════
# Word 教案生成器
# ═══════════════════════════════════════════

class DocxGenerator:
    """
    Word 教案自动生成器。

    Example:
        gen = DocxGenerator(config_path="config.yaml")
        result = gen.generate_from_text(
            topic="中心法则 — 转录过程",
            teaching_goal="学生能描述转录的基本步骤",
            output_dir="outputs/docx",
        )
        print(result["docx_path"])
    """

    def __init__(self, config_path: str = "config.yaml"):
        self.config_path = config_path
        self._generator = None

    @property
    def generator(self):
        if self._generator is None:
            from .config import load_config
            from .generator import GeneratorConfig, QwenGenerator
            cfg = load_config(self.config_path)
            g_cfg = cfg.get("generator", {}) or {}
            self._generator = QwenGenerator(GeneratorConfig(
                mode=str(g_cfg.get("mode", "api")),
                api_model=str(g_cfg.get("api_model", "qwen-plus-latest")),
                api_base_url_env=str(g_cfg.get("api_base_url_env", "OPENAI_BASE_URL")),
                api_key_env=str(g_cfg.get("api_key_env", "OPENAI_API_KEY")),
                timeout_sec=float(g_cfg.get("timeout_sec", 120.0)),
                local_model_path=str(g_cfg.get("local_model_path", "")),
                device=str(g_cfg.get("device", "cpu")),
                max_new_tokens=int(g_cfg.get("max_new_tokens", 4096)),
                temperature=float(g_cfg.get("temperature", 0.3)),
            ))
        return self._generator

    def generate_from_outline(
        self,
        outline: Dict[str, Any],
        output_dir: str = "outputs/docx",
    ) -> Dict[str, Any]:
        """
        根据结构化大纲生成Word教案。

        outline 格式:
        {
          "title": "教案标题",
          "subject": "学科",
          "grade": "年级",
          "duration": "45分钟",
          "teaching_goal": ["目标1", "目标2"],
          "key_points": ["重点1"],
          "difficulties": ["难点1"],
          "teaching_method": ["讲授法", "讨论法"],
          "teaching_process": [
            {"stage": "导入", "duration": "5分钟", "content": "...", "activity": "..."},
            {"stage": "新授", "duration": "25分钟", "content": "...", "activity": "..."},
          ],
          "homework": ["作业1", "作业2"],
          "reflection": "教学反思"
        }
        """
        t0 = time.time()
        doc = Document()

        title = outline.get("title", "教学教案")

        # === 标题 ===
        _add_heading(doc, title, level=0)
        _add_para(doc, "", font_size=6)

        # === 基本信息表 ===
        info_rows = []
        if outline.get("subject"):
            info_rows.append(["学科", outline["subject"]])
        if outline.get("grade"):
            info_rows.append(["年级", outline["grade"]])
        if outline.get("duration"):
            info_rows.append(["课时", outline["duration"]])
        if outline.get("textbook"):
            info_rows.append(["教材", outline["textbook"]])

        if info_rows:
            _add_table(doc, ["项目", "内容"], info_rows)

        # === 教学目标 ===
        goals = outline.get("teaching_goal", [])
        if goals:
            _add_heading(doc, "一、教学目标", level=1)
            if isinstance(goals, str):
                goals = [goals]
            for i, g in enumerate(goals, 1):
                _add_para(doc, f"{i}. {g}")

        # === 重点难点 ===
        key_points = outline.get("key_points", [])
        difficulties = outline.get("difficulties", [])
        if key_points or difficulties:
            _add_heading(doc, "二、重点与难点", level=1)
            if key_points:
                _add_para(doc, "【教学重点】", bold=True)
                for kp in key_points:
                    _add_para(doc, f"  • {kp}")
            if difficulties:
                _add_para(doc, "【教学难点】", bold=True)
                for d in difficulties:
                    _add_para(doc, f"  • {d}")

        # === 教学方法 ===
        methods = outline.get("teaching_method", [])
        if methods:
            _add_heading(doc, "三、教学方法", level=1)
            _add_para(doc, "、".join(methods))

        # === 教学过程 ===
        process = outline.get("teaching_process", [])
        if process:
            _add_heading(doc, "四、教学过程", level=1)
            headers = ["教学环节", "时间", "教学内容", "教学活动"]
            rows = []
            for step in process:
                rows.append([
                    step.get("stage", ""),
                    step.get("duration", ""),
                    step.get("content", ""),
                    step.get("activity", ""),
                ])
            _add_table(doc, headers, rows)

        # === 课堂活动设计 ===
        activities = outline.get("activities", [])
        if activities:
            _add_heading(doc, "五、课堂活动设计", level=1)
            for i, act in enumerate(activities, 1):
                if isinstance(act, dict):
                    _add_para(doc, f"活动{i}: {act.get('name', '')}", bold=True)
                    _add_para(doc, act.get("description", ""))
                else:
                    _add_para(doc, f"{i}. {act}")

        # === 课后作业 ===
        homework = outline.get("homework", [])
        if homework:
            _add_heading(doc, "六、课后作业", level=1)
            for i, hw in enumerate(homework, 1):
                _add_para(doc, f"{i}. {hw}")

        # === 教学反思 ===
        reflection = outline.get("reflection", "")
        if reflection:
            _add_heading(doc, "七、教学反思", level=1)
            _add_para(doc, reflection)

        # === 板书设计 ===
        board = outline.get("board_design", "")
        if board:
            _add_heading(doc, "八、板书设计", level=1)
            _add_para(doc, board)

        # 保存
        Path(output_dir).mkdir(parents=True, exist_ok=True)
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_title = re.sub(r'[^\w\u4e00-\u9fff]', '_', title)[:30]
        filename = f"docx_{safe_title}_{ts}.docx"
        filepath = str(Path(output_dir) / filename)
        doc.save(filepath)

        gen_time = round(time.time() - t0, 2)
        return {
            "docx_path": filepath,
            "title": title,
            "generation_time": gen_time,
        }

    def generate_from_text(
        self,
        topic: str,
        teaching_goal: str = "",
        key_points: str = "",
        duration: str = "45分钟",
        context: str = "",
        output_dir: str = "outputs/docx",
    ) -> Dict[str, Any]:
        """从自然语言需求生成教案（先调LLM生成大纲，再渲染）"""
        prompt = f"""你是一位经验丰富的教学设计专家。请根据以下需求，生成一份详细的教学教案。

【课题】{topic}
【教学目标】{teaching_goal or '请根据课题自行设定'}
【重点内容】{key_points or '请根据课题自行设定'}
【课时】{duration}
【参考内容】
{context[:3000] if context else '无'}

请严格按以下JSON格式输出，不要输出多余文字：
{{
  "title": "教案标题",
  "subject": "学科",
  "grade": "适用年级",
  "duration": "{duration}",
  "teaching_goal": ["教学目标1", "教学目标2", "教学目标3"],
  "key_points": ["教学重点1", "教学重点2"],
  "difficulties": ["教学难点1"],
  "teaching_method": ["讲授法", "讨论法", "演示法"],
  "teaching_process": [
    {{"stage": "导入新课", "duration": "5分钟", "content": "教学内容", "activity": "教学活动描述"}},
    {{"stage": "讲授新知", "duration": "20分钟", "content": "教学内容", "activity": "教学活动描述"}},
    {{"stage": "巩固练习", "duration": "10分钟", "content": "教学内容", "activity": "教学活动描述"}},
    {{"stage": "课堂小结", "duration": "5分钟", "content": "教学内容", "activity": "教学活动描述"}},
    {{"stage": "布置作业", "duration": "5分钟", "content": "教学内容", "activity": "教学活动描述"}}
  ],
  "homework": ["课后作业1", "课后作业2"],
  "activities": [
    {{"name": "课堂活动名称", "description": "活动详细描述：目的、形式、时长、预期效果"}},
    {{"name": "第二个活动", "description": "描述"}}
  ],
  "board_design": "板书设计描述",
  "reflection": "教学反思要点"
}}"""

        raw = self.generator.generate(prompt)
        outline = self._parse_json(raw)

        if not outline.get("title"):
            outline["title"] = topic

        return self.generate_from_outline(outline, output_dir)

    @staticmethod
    def _parse_json(raw: str) -> Dict[str, Any]:
        text = raw.strip()
        text = re.sub(r'^```(?:json)?\s*', '', text, flags=re.MULTILINE)
        text = re.sub(r'```\s*$', '', text, flags=re.MULTILINE)
        text = text.strip()
        try:
            return json.loads(text)
        except json.JSONDecodeError:
            pass
        match = re.search(r'\{[\s\S]*\}', text)
        if match:
            candidate = re.sub(r',\s*([}\]])', r'\1', match.group(0))
            try:
                return json.loads(candidate)
            except json.JSONDecodeError:
                pass
        return {}
